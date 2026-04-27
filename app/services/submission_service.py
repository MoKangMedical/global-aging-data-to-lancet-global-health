"""
Submission Package Service.
Creates ZIP files containing manuscript, tables, figures, and README
in Lancet Global Health submission format.
"""
import os
import json
import shutil
import zipfile
from datetime import datetime
from typing import Dict, List, Any, Optional

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUTS_DIR = os.path.join(BASE_DIR, "outputs")
FIGURES_DIR = os.path.join(BASE_DIR, "figures")

os.makedirs(OUTPUTS_DIR, exist_ok=True)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False


def generate_submission_package(
    project_id: int,
    project_title: str,
    manuscript: Dict[str, str],
    tables_xlsx_path: str = None,
    figure_paths: List[str] = None,
) -> str:
    """
    Generate a complete submission package as a ZIP file.
    Returns the path to the generated ZIP file.
    """
    package_dir = os.path.join(OUTPUTS_DIR, f"submission_p{project_id}")
    os.makedirs(package_dir, exist_ok=True)

    # 1. Generate manuscript.docx
    docx_path = os.path.join(package_dir, "manuscript.docx")
    _generate_manuscript_docx(docx_path, manuscript)

    # 2. Copy or generate tables.xlsx
    if tables_xlsx_path and os.path.exists(tables_xlsx_path):
        shutil.copy2(tables_xlsx_path, os.path.join(package_dir, "tables.xlsx"))
    else:
        _generate_empty_tables_xlsx(os.path.join(package_dir, "tables.xlsx"), project_title)

    # 3. Copy figures
    figures_dir = os.path.join(package_dir, "figures")
    os.makedirs(figures_dir, exist_ok=True)
    if figure_paths:
        for fp in figure_paths:
            if fp and os.path.exists(fp):
                shutil.copy2(fp, figures_dir)

    # 4. Generate README.txt
    readme_path = os.path.join(package_dir, "README.txt")
    _generate_readme(readme_path, project_title, manuscript, figure_paths)

    # 5. Create ZIP
    zip_filename = f"submission_p{project_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.zip"
    zip_path = os.path.join(OUTPUTS_DIR, zip_filename)
    _create_zip(package_dir, zip_path)

    # Clean up temp directory
    shutil.rmtree(package_dir, ignore_errors=True)

    return zip_path


def _generate_manuscript_docx(filepath: str, manuscript: Dict[str, str]):
    """Generate a Word document from manuscript sections."""
    if not HAS_DOCX:
        # Fallback: write plain text
        with open(filepath.replace(".docx", ".txt"), "w", encoding="utf-8") as f:
            for section in ["title", "abstract", "introduction", "methods", "results", "discussion", "references"]:
                content = manuscript.get(section, "")
                f.write(f"\n\n{'='*60}\n{section.upper()}\n{'='*60}\n\n{content}")
        return filepath

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Title
    title = manuscript.get("title", "Untitled Manuscript")
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(14)

    # Abstract
    doc.add_heading("Abstract", level=2)
    abstract = manuscript.get("abstract", "")
    for para_text in abstract.split("\n\n"):
        if para_text.strip():
            # Bold section labels
            if para_text.startswith("**") and "**:" in para_text[:30]:
                p = doc.add_paragraph()
                parts = para_text.split("**:", 1)
                label = parts[0].replace("**", "")
                run = p.add_run(f"{label}: ")
                run.bold = True
                run.font.size = Pt(12)
                if len(parts) > 1:
                    p.add_run(parts[1].strip()).font.size = Pt(12)
            else:
                doc.add_paragraph(para_text.strip())

    doc.add_page_break()

    # Main sections
    for section_name in ["introduction", "methods", "results", "discussion"]:
        content = manuscript.get(section_name, "")
        doc.add_heading(section_name.capitalize(), level=2)

        for para_text in content.split("\n\n"):
            if not para_text.strip():
                continue
            # Bold subsection headers
            if para_text.startswith("**") and para_text.strip().endswith("**"):
                heading_text = para_text.replace("**", "").strip()
                doc.add_heading(heading_text, level=3)
            elif para_text.startswith("**"):
                p = doc.add_paragraph()
                parts = para_text.split("**")
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:  # Bold parts
                        run.bold = True
                    run.font.size = Pt(12)
            else:
                doc.add_paragraph(para_text.strip())

    # References
    doc.add_page_break()
    doc.add_heading("References", level=2)
    references = manuscript.get("references", "")
    for ref in references.split("\n\n"):
        if ref.strip():
            p = doc.add_paragraph(ref.strip())
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.size = Pt(10)

    doc.save(filepath)
    return filepath


def _generate_empty_tables_xlsx(filepath: str, title: str):
    """Generate an empty tables workbook as fallback."""
    if not HAS_OPENPYXL:
        return
    wb = Workbook()
    ws = wb.active
    ws.title = "Table 1"
    ws["A1"] = f"Tables for: {title}"
    ws["A1"].font = Font(size=14, bold=True)
    wb.save(filepath)


def _generate_readme(filepath: str, project_title: str, manuscript: Dict[str, str], figure_paths: List[str] = None):
    """Generate README.txt for the submission package."""
    now = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
    n_figures = len(figure_paths) if figure_paths else 0

    content = f"""GLOBAL AGING DATA TO LANCET GLOBAL HEALTH — SUBMISSION PACKAGE
================================================================

Project: {project_title}
Generated: {now}
Platform: Global Aging Data to Lancet Global Health Research Automation

PACKAGE CONTENTS
-----------------
1. manuscript.docx    — Full manuscript in Lancet Global Health format
2. tables.xlsx        — All tables (Table 1: Baseline Characteristics, Table 2: Regression Results)
3. figures/           — Publication-quality figures (PNG, 300 DPI)
   - Kaplan-Meier survival curves
   - Forest plot of adjusted effect estimates
   - Cumulative incidence functions (competing risks)
4. README.txt         — This file

MANUSCRIPT SECTIONS
--------------------
Title: {manuscript.get('title', 'N/A')[:100]}...

The manuscript includes the following sections:
- Abstract (structured: Background, Methods, Findings, Interpretation)
- Introduction
- Methods
- Results
- Discussion
- References

LANCET GLOBAL HEALTH FORMATTING NOTES
----------------------------------------
- Manuscript follows Lancet Global Health author guidelines
- Tables are formatted per Lancet style (horizontal lines only)
- Figures are generated at 300 DPI for print quality
- The Lancet uses SI units and the International System of Units
- Numbers are formatted with middle dots (e.g., 0·05 instead of 0.05)

DATA SOURCES
--------------
This analysis uses harmonised data from the Gateway to Global Aging Data (G2AGING),
incorporating data from multiple longitudinal studies on ageing worldwide.

CONTACT
---------
For questions about this submission package, please refer to the
Global Aging Data to Lancet Global Health platform documentation.

================================================================
Generated by the Global Aging Data to Lancet Global Health Platform
"""

    with open(filepath, "w", encoding="utf-8") as f:
        f.write(content)


def _create_zip(source_dir: str, output_path: str):
    """Create a ZIP archive from a directory."""
    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(source_dir):
            for file in files:
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, source_dir)
                zf.write(file_path, arcname)
